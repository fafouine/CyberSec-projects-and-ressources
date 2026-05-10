"""
©AngelaMos | 2026
office.py
"""


import structlog

from dlp_scanner.models import Location, TextChunk


log = structlog.get_logger()

DOCX_EXTENSIONS: frozenset[str] = frozenset({".docx"})
XLSX_EXTENSIONS: frozenset[str] = frozenset({".xlsx"})
XLS_EXTENSIONS: frozenset[str] = frozenset({".xls"})
OFFICE_EXTENSIONS: frozenset[str] = (
    DOCX_EXTENSIONS | XLSX_EXTENSIONS | XLS_EXTENSIONS
)


class DocxExtractor:
    """
    Extracts text from DOCX files
    """
    @property
    def supported_extensions(self) -> frozenset[str]:
        """
        File extensions this extractor handles
        """
        return DOCX_EXTENSIONS

    def extract(self, path: str) -> list[TextChunk]:
        """
        Extract text from paragraphs, tables, and metadata
        """
        from docx import Document

        chunks: list[TextChunk] = []

        try:
            doc = Document(path)
            paragraphs: list[str] = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    cells = [
                        cell.text
                        for cell in row.cells
                        if cell.text.strip()
                    ]
                    if cells:
                        paragraphs.append(" | ".join(cells))

            if doc.core_properties.author:
                paragraphs.append(f"Author: {doc.core_properties.author}")
            if doc.core_properties.title:
                paragraphs.append(f"Title: {doc.core_properties.title}")

            if paragraphs:
                chunks.append(
                    TextChunk(
                        text = "\n".join(paragraphs),
                        location = Location(
                            source_type = "file",
                            uri = path,
                        ),
                    )
                )

        except Exception:
            log.warning("docx_extract_failed", path = path)

        return chunks


class XlsxExtractor:
    """
    Extracts text from XLSX files using openpyxl
    """
    @property
    def supported_extensions(self) -> frozenset[str]:
        """
        File extensions this extractor handles
        """
        return XLSX_EXTENSIONS

    def extract(self, path: str) -> list[TextChunk]:
        """
        Extract text from all sheets and cells
        """
        from openpyxl import load_workbook

        chunks: list[TextChunk] = []

        try:
            wb = load_workbook(
                path,
                read_only = True,
                data_only = True,
            )
            for sheet in wb.worksheets:
                rows: list[str] = []
                for row in sheet.iter_rows(values_only = True):
                    cell_values = [str(c) for c in row if c is not None]
                    if cell_values:
                        rows.append(" | ".join(cell_values))

                if rows:
                    chunks.append(
                        TextChunk(
                            text = "\n".join(rows),
                            location = Location(
                                source_type = "file",
                                uri = path,
                                sheet_name = sheet.title,
                            ),
                        )
                    )
            wb.close()
        except Exception:
            log.warning("xlsx_extract_failed", path = path)

        return chunks


class XlsExtractor:
    """
    Extracts text from legacy XLS files using xlrd
    """
    @property
    def supported_extensions(self) -> frozenset[str]:
        """
        File extensions this extractor handles
        """
        return XLS_EXTENSIONS

    def extract(self, path: str) -> list[TextChunk]:
        """
        Extract text from legacy Excel workbooks
        """
        import xlrd

        chunks: list[TextChunk] = []

        try:
            wb = xlrd.open_workbook(path)
            for sheet in wb.sheets():
                rows: list[str] = []
                for row_idx in range(sheet.nrows):
                    cell_values = [
                        str(sheet.cell_value(row_idx,
                                             col))
                        for col in range(sheet.ncols)
                        if sheet.cell_value(row_idx, col)
                    ]
                    if cell_values:
                        rows.append(" | ".join(cell_values))

                if rows:
                    chunks.append(
                        TextChunk(
                            text = "\n".join(rows),
                            location = Location(
                                source_type = "file",
                                uri = path,
                                sheet_name = sheet.name,
                            ),
                        )
                    )
        except Exception:
            log.warning("xls_extract_failed", path = path)

        return chunks
